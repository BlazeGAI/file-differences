import streamlit as st
from docx import Document
from deepdiff import DeepDiff
import pandas as pd
import webcolors
import os
from collections import defaultdict

st.title("Word Document Checker 📄")

# Upload Master and Student Documents
st.header("Upload Documents for Comparison")
word_file_master = st.file_uploader("Upload Master Document (Correct Version)", type=["docx"], key="master_doc")
word_file_student = st.file_uploader("Upload Student Document", type=["docx"], key="student_doc")

def closest_color(hex_code):
    """Convert hex color codes to the closest known color name."""
    if not hex_code:
        return None
    hex_code = f"#{hex_code}" if not hex_code.startswith("#") else hex_code
    try:
        return webcolors.hex_to_name(hex_code)
    except ValueError:
        return hex_code  # Return hex code if no name found

def get_paragraph_style_info(paragraph):
    """Extract detailed style information from a paragraph."""
    style_info = defaultdict(list)
    
    # Track position of each run for accurate comparison
    position = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        
        # Basic run properties
        if run.bold:
            style_info["bold"].append((position, position + run_length))
        if run.italic:
            style_info["italic"].append((position, position + run_length))
        if run.underline:
            style_info["underline"].append((position, position + run_length))
            
        # Font properties
        rpr = run._element.find("w:rPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        if rpr is not None:
            # Font size
            sz = rpr.find("w:sz", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if sz is not None:
                size = int(sz.attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"]) / 2
                style_info["font_size"].append((position, position + run_length, size))
            
            # Font color
            color = rpr.find("w:color", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if color is not None:
                color_val = closest_color(color.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
                style_info["font_color"].append((position, position + run_length, color_val))
            
            # Background color
            highlight = rpr.find("w:highlight", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if highlight is not None:
                highlight_val = closest_color(highlight.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
                style_info["background_color"].append((position, position + run_length, highlight_val))
        
        position += run_length
    
    return style_info

def extract_text_with_styles(doc):
    """Extracts text and formatting with improved accuracy."""
    content = []
    tables = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_info = get_paragraph_style_info(para)
        
        content.append({
            "text": text,
            "style_info": style_info,
            "heading_level": para.style.name if para.style.name.startswith("Heading") else None,
            "alignment": para.alignment,
            "style_name": para.style.name
        })
    
    # Process tables
    for table in doc.tables:
        table_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                cell_content = []
                for para in cell.paragraphs:
                    if para.text.strip():
                        style_info = get_paragraph_style_info(para)
                        cell_content.append({
                            "text": para.text.strip(),
                            "style_info": style_info,
                            "alignment": para.alignment
                        })
                
                # Get cell properties
                tc = cell._tc
                tc_pr = tc.find("w:tcPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                cell_properties = {
                    "content": cell_content,
                    "background_color": None,
                    "borders": {
                        "top": False,
                        "bottom": False,
                        "left": False,
                        "right": False
                    },
                    "position": (row_idx, col_idx)
                }
                
                if tc_pr is not None:
                    # Background color
                    shd = tc_pr.find("w:shd", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if shd is not None:
                        cell_properties["background_color"] = closest_color(
                            shd.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                        )
                    
                    # Borders
                    borders = tc_pr.find("w:tcBorders", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if borders is not None:
                        for border in ["top", "bottom", "left", "right"]:
                            border_elem = borders.find(f"w:{border}", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                            cell_properties["borders"][border] = border_elem is not None
                
                row_data.append(cell_properties)
            table_data.append(row_data)
        tables.append(table_data)
    
    return content, tables

def compare_documents(master_doc, student_doc):
    """Compare documents with improved difference detection."""
    master_content, master_tables = extract_text_with_styles(master_doc)
    student_content, student_tables = extract_text_with_styles(student_doc)
    
    differences = []
    
    # Compare paragraphs
    for idx, (master_para, student_para) in enumerate(zip(master_content, student_content)):
        # Text differences
        if master_para["text"] != student_para["text"]:
            differences.append({
                "type": "Text",
                "location": f"Paragraph {idx + 1}",
                "student": student_para["text"],
                "master": master_para["text"]
            })
        
        # Style differences
        for style_type, master_styles in master_para["style_info"].items():
            student_styles = student_para["style_info"][style_type]
            if master_styles != student_styles:
                differences.append({
                    "type": style_type.replace("_", " ").title(),
                    "location": f"Paragraph {idx + 1}",
                    "student": str(student_styles),
                    "master": str(master_styles)
                })
        
        # Heading level differences
        if master_para["heading_level"] != student_para["heading_level"]:
            differences.append({
                "type": "Heading Level",
                "location": f"Paragraph {idx + 1}",
                "student": student_para["heading_level"],
                "master": master_para["heading_level"]
            })
    
    # Compare tables
    for table_idx, (master_table, student_table) in enumerate(zip(master_tables, student_tables)):
        for row_idx, (master_row, student_row) in enumerate(zip(master_table, student_table)):
            for col_idx, (master_cell, student_cell) in enumerate(zip(master_row, student_row)):
                location = f"Table {table_idx + 1}, Row {row_idx + 1}, Column {col_idx + 1}"
                
                # Compare cell content
                for m_content, s_content in zip(master_cell["content"], student_cell["content"]):
                    if m_content["text"] != s_content["text"]:
                        differences.append({
                            "type": "Table Cell Text",
                            "location": location,
                            "student": s_content["text"],
                            "master": m_content["text"]
                        })
                
                # Compare cell properties
                if master_cell["background_color"] != student_cell["background_color"]:
                    differences.append({
                        "type": "Table Cell Background",
                        "location": location,
                        "student": student_cell["background_color"],
                        "master": master_cell["background_color"]
                    })
                
                # Compare borders
                for border_type, master_border in master_cell["borders"].items():
                    student_border = student_cell["borders"][border_type]
                    if master_border != student_border:
                        differences.append({
                            "type": f"Table Cell {border_type.title()} Border",
                            "location": location,
                            "student": "Present" if student_border else "Missing",
                            "master": "Present" if master_border else "Missing"
                        })
    
    return pd.DataFrame(differences)

if word_file_master and word_file_student:
    try:
        with open("master.docx", "wb") as f1, open("student.docx", "wb") as f2:
            f1.write(word_file_master.getbuffer())
            f2.write(word_file_student.getbuffer())

        master_doc = Document("master.docx")
        student_doc = Document("student.docx")

        st.subheader("Comparison Results")
        
        differences_df = compare_documents(master_doc, student_doc)
        
        if not differences_df.empty:
            st.write("### Differences Found:")
            # Add filtering options
            diff_types = ["All"] + list(differences_df["type"].unique())
            selected_type = st.selectbox("Filter by difference type:", diff_types)
            
            if selected_type != "All":
                filtered_df = differences_df[differences_df["type"] == selected_type]
            else:
                filtered_df = differences_df
            
            # Display differences with better formatting
            st.dataframe(
                filtered_df,
                column_config={
                    "type": "Difference Type",
                    "location": "Location",
                    "student": "Student Version",
                    "master": "Master Version"
                },
                hide_index=True
            )
            
            # Summary statistics
            st.write("### Summary Statistics")
            type_counts = differences_df["type"].value_counts()
            st.bar_chart(type_counts)
        else:
            st.success("✅ No differences found. The student document matches the master.")
    
    except Exception as e:
        st.error(f"An error occurred while comparing documents: {str(e)}")
    finally:
        # Clean up temporary files
        for filename in ["master.docx", "student.docx"]:
            if os.path.exists(filename):
                os.remove(filename)
